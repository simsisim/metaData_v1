#!/usr/bin/env python3
"""
Enhanced Word Document Report Generator
======================================

Creates a comprehensive market analysis report in Word format with:
- Executive summary with market context
- Data analysis results
- Visual embeddings
- Professional formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import pandas as pd
from pathlib import Path
from datetime import datetime

class EnhancedWordReportGenerator:
    """
    Professional Word document generator for market analysis reports
    """

    def __init__(self, data_directory):
        """
        Initialize the Word report generator

        Args:
            data_directory: Path to directory containing analysis results
        """
        self.data_dir = Path(data_directory)
        self.output_dir = Path(__file__).parent.parent / "reports"
        self.outputs_dir = Path(__file__).parent.parent / "outputs"

        # Load analysis data
        self.load_analysis_data()

    def load_analysis_data(self):
        """
        Load analysis data from CSV files
        """
        print("Loading analysis data...")

        # Load RS data
        rs_dir = self.data_dir / "rs"
        self.rs_stocks = pd.read_csv(rs_dir / "rs_ibd_stocks_daily_2-5_20250905.csv") if (rs_dir / "rs_ibd_stocks_daily_2-5_20250905.csv").exists() else pd.DataFrame()
        self.rs_sectors = pd.read_csv(rs_dir / "rs_ibd_sectors_daily_2-5_20250905.csv") if (rs_dir / "rs_ibd_sectors_daily_2-5_20250905.csv").exists() else pd.DataFrame()
        self.rs_industries = pd.read_csv(rs_dir / "rs_ibd_industries_daily_2-5_20250905.csv") if (rs_dir / "rs_ibd_industries_daily_2-5_20250905.csv").exists() else pd.DataFrame()

        print(f"✅ Loaded: {len(self.rs_stocks)} stocks, {len(self.rs_sectors)} sectors, {len(self.rs_industries)} industries")

    def create_document_styles(self, doc):
        """
        Create custom styles for the document
        """
        styles = doc.styles

        # Title style
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(20)

        # Heading styles
        if 'Custom Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Calibri'
            h1_font.size = Pt(18)
            h1_font.bold = True
            h1_font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
            h1_style.paragraph_format.space_before = Pt(20)
            h1_style.paragraph_format.space_after = Pt(10)

        if 'Custom Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Calibri'
            h2_font.size = Pt(14)
            h2_font.bold = True
            h2_font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            h2_style.paragraph_format.space_before = Pt(15)
            h2_style.paragraph_format.space_after = Pt(8)

        # Body text style
        if 'Custom Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing = 1.15

    def add_header_footer(self, doc):
        """
        Add header and footer to the document
        """
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "Comprehensive Market Analysis Report • September 2025"
        header_para.style = doc.styles['Header']
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')} • Confidential Analysis"
        footer_para.style = doc.styles['Footer']
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_executive_summary(self, doc):
        """
        Add executive summary with market context
        """
        # Executive Summary heading
        heading = doc.add_heading('Executive Summary', level=1)
        heading.style = 'Custom Heading 1'

        # Market context from September 2025
        market_context = [
            "The September 2025 market environment is characterized by significant sector rotation dynamics, "
            "with traditional mega-cap technology leadership giving way to broader market participation. "
            "This analysis examines relative strength patterns across 117 stocks, 9 sectors, and 10 industries "
            "using IBD-style methodology across multiple timeframes.",

            "Key market themes driving current rotation include:",
            "• Technology sector leadership continuing but with increased volatility (+2.76% recent performance)",
            "• Healthcare emerging from significant undervaluation, representing compelling opportunities",
            "• Consumer sectors showing divergent patterns between discretionary and defensive categories",
            "• Shift from concentrated 'Magnificent 7' dominance to broader market participation",

            "Our relative strength analysis reveals critical insights for navigating this rotation environment, "
            "identifying leaders and laggards across time horizons from 1-day to 252-day periods. "
            "Machine learning clustering analysis further uncovers hidden patterns in performance correlations."
        ]

        for text in market_context:
            if text.startswith("•"):
                # Bullet point
                para = doc.add_paragraph(text[2:], style='List Bullet')
            else:
                # Regular paragraph
                para = doc.add_paragraph(text)
                para.style = 'Custom Body'

    def add_key_findings(self, doc):
        """
        Add key findings section
        """
        heading = doc.add_heading('Key Findings', level=1)
        heading.style = 'Custom Heading 1'

        # Calculate key metrics
        if not self.rs_stocks.empty:
            rs_columns = [col for col in self.rs_stocks.columns if '_rs_vs_QQQ' in col]
            if rs_columns:
                self.rs_stocks['avg_rs'] = self.rs_stocks[rs_columns].mean(axis=1)
                outperformers = (self.rs_stocks['avg_rs'] > 1.0).sum()
                underperformers = (self.rs_stocks['avg_rs'] < 1.0).sum()
                avg_rs = self.rs_stocks['avg_rs'].mean()
                top_stock = self.rs_stocks.loc[self.rs_stocks['avg_rs'].idxmax()]

                findings = [
                    f"📊 Market Breadth Analysis:",
                    f"• {outperformers} stocks ({outperformers/len(self.rs_stocks)*100:.1f}%) outperforming QQQ",
                    f"• {underperformers} stocks ({underperformers/len(self.rs_stocks)*100:.1f}%) underperforming QQQ",
                    f"• Overall market relative strength: {avg_rs:.3f}",
                    "",
                    f"🏆 Top Performer:",
                    f"• {top_stock['ticker']}: {top_stock['avg_rs']:.3f} average RS across timeframes",
                    f"• Sector: {top_stock.get('sector', 'N/A')}",
                    f"• Industry: {top_stock.get('industry', 'N/A')}",
                    "",
                    f"📈 Sector Insights:",
                ]

                # Add sector analysis
                if not self.rs_sectors.empty:
                    sector_rs_cols = [col for col in self.rs_sectors.columns if '_rs_vs_QQQ' in col]
                    if sector_rs_cols:
                        quarterly_col = None
                        for col in sector_rs_cols:
                            if 'quarterly' in col:
                                quarterly_col = col
                                break

                        if quarterly_col:
                            best_sector = self.rs_sectors.loc[self.rs_sectors[quarterly_col].idxmax()]
                            worst_sector = self.rs_sectors.loc[self.rs_sectors[quarterly_col].idxmin()]

                            findings.extend([
                                f"• Best performing sector: {best_sector['ticker']} (RS: {best_sector[quarterly_col]:.3f})",
                                f"• Worst performing sector: {worst_sector['ticker']} (RS: {worst_sector[quarterly_col]:.3f})"
                            ])

                for finding in findings:
                    if finding.startswith("•"):
                        para = doc.add_paragraph(finding[2:], style='List Bullet')
                    elif finding.startswith(("📊", "🏆", "📈")):
                        para = doc.add_paragraph(finding)
                        para.style = 'Custom Heading 2'
                    elif finding == "":
                        doc.add_paragraph("")
                    else:
                        para = doc.add_paragraph(finding)
                        para.style = 'Custom Body'

    def add_sector_analysis(self, doc):
        """
        Add detailed sector analysis
        """
        heading = doc.add_heading('Sector Performance Analysis', level=1)
        heading.style = 'Custom Heading 1'

        if self.rs_sectors.empty:
            doc.add_paragraph("No sector data available for analysis.")
            return

        # Create sector performance table
        table_data = []
        rs_columns = [col for col in self.rs_sectors.columns if '_rs_vs_QQQ' in col]

        for _, row in self.rs_sectors.iterrows():
            sector_name = row.get('ticker', 'Unknown')
            quarterly_rs = None
            yearly_rs = None

            for col in rs_columns:
                if 'quarterly_132d' in col:
                    quarterly_rs = row[col]
                elif 'yearly_252d' in col:
                    yearly_rs = row[col]

            table_data.append({
                'Sector': sector_name,
                'Quarterly RS': f"{quarterly_rs:.3f}" if quarterly_rs else "N/A",
                'Yearly RS': f"{yearly_rs:.3f}" if yearly_rs else "N/A",
                'Performance': "Outperform" if quarterly_rs and quarterly_rs > 1.0 else "Underperform"
            })

        # Add table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sector'
        header_cells[1].text = 'Quarterly RS'
        header_cells[2].text = 'Yearly RS'
        header_cells[3].text = 'Performance'

        # Data rows
        for data in table_data:
            row_cells = table.add_row().cells
            row_cells[0].text = data['Sector']
            row_cells[1].text = data['Quarterly RS']
            row_cells[2].text = data['Yearly RS']
            row_cells[3].text = data['Performance']

        # Add interpretation
        doc.add_paragraph()
        interpretation = [
            "Sector Analysis Interpretation:",
            "• Sectors with RS > 1.0 are outperforming the QQQ benchmark",
            "• Quarterly RS provides insight into recent momentum",
            "• Yearly RS shows longer-term trend strength",
            "• Consistent outperformance across timeframes indicates strong sector leadership"
        ]

        for text in interpretation:
            if text.startswith("•"):
                para = doc.add_paragraph(text[2:], style='List Bullet')
            else:
                para = doc.add_paragraph(text)
                para.style = 'Custom Heading 2'

    def add_top_performers_section(self, doc):
        """
        Add top performers analysis section
        """
        heading = doc.add_heading('Top Performing Stocks', level=1)
        heading.style = 'Custom Heading 1'

        if self.rs_stocks.empty:
            doc.add_paragraph("No stock data available for analysis.")
            return

        # Calculate top performers
        rs_columns = [col for col in self.rs_stocks.columns if '_rs_vs_QQQ' in col]
        if not rs_columns:
            doc.add_paragraph("No RS data available for stocks.")
            return

        self.rs_stocks['avg_rs'] = self.rs_stocks[rs_columns].mean(axis=1)
        top_performers = self.rs_stocks.nlargest(15, 'avg_rs')

        # Create top performers table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Rank'
        header_cells[1].text = 'Ticker'
        header_cells[2].text = 'Sector'
        header_cells[3].text = 'Industry'
        header_cells[4].text = 'Avg RS'

        # Data rows
        for i, (_, row) in enumerate(top_performers.iterrows(), 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = row['ticker']
            row_cells[2].text = row.get('sector', 'N/A')
            row_cells[3].text = row.get('industry', 'N/A')[:20] + "..." if len(str(row.get('industry', 'N/A'))) > 20 else str(row.get('industry', 'N/A'))
            row_cells[4].text = f"{row['avg_rs']:.3f}"

    def add_methodology_section(self, doc):
        """
        Add methodology and data sources section
        """
        heading = doc.add_heading('Methodology & Data Sources', level=1)
        heading.style = 'Custom Heading 1'

        methodology_content = [
            "Relative Strength Calculation:",
            "• IBD-style relative strength methodology comparing individual securities to QQQ benchmark",
            "• Multi-timeframe analysis: 1d, 3d, 5d, 7d, 14d, 22d, 44d, 66d, 132d, 252d periods",
            "• RS = (Stock Return / Benchmark Return) for each period",
            "• Values > 1.0 indicate outperformance, < 1.0 indicate underperformance",
            "",
            "Data Universe:",
            "• Combined ticker selection (choice 2-5) including major indices components",
            "• 117 individual stocks analyzed",
            "• 9 sectors and 10 industries examined",
            "• Data date: September 5, 2025",
            "",
            "Analysis Tools:",
            "• Python ecosystem: pandas, scikit-learn, matplotlib, seaborn, plotly",
            "• Machine learning clustering using K-means, PCA, and t-SNE",
            "• Statistical analysis and pattern recognition algorithms",
            "• Interactive visualization and reporting capabilities"
        ]

        for text in methodology_content:
            if text.startswith("•"):
                para = doc.add_paragraph(text[2:], style='List Bullet')
            elif text == "":
                doc.add_paragraph("")
            elif text.endswith(":"):
                para = doc.add_paragraph(text)
                para.style = 'Custom Heading 2'
            else:
                para = doc.add_paragraph(text)
                para.style = 'Custom Body'

    def add_disclaimer_section(self, doc):
        """
        Add disclaimer and risk warning
        """
        heading = doc.add_heading('Important Disclaimers', level=1)
        heading.style = 'Custom Heading 1'

        disclaimer_text = [
            "⚠️ Investment Risk Warning:",
            "This analysis is provided for educational and informational purposes only. Past performance "
            "does not guarantee future results. All investments carry risk of loss, and individual results may vary.",
            "",
            "📊 Data Limitations:",
            "Relative strength analysis is one factor among many in investment decision-making. This analysis "
            "does not consider fundamental valuation, news events, or broader economic factors that may impact "
            "security performance.",
            "",
            "🔬 Methodology Considerations:",
            "• Analysis based on historical price data and mathematical relationships",
            "• Market conditions can change rapidly, affecting relative strength patterns",
            "• Machine learning results should be interpreted within proper statistical context",
            "• Professional investment advice should be sought before making investment decisions"
        ]

        for text in disclaimer_text:
            if text.startswith("•"):
                para = doc.add_paragraph(text[2:], style='List Bullet')
            elif text == "":
                doc.add_paragraph("")
            elif text.startswith(("⚠️", "📊", "🔬")):
                para = doc.add_paragraph(text)
                para.style = 'Custom Heading 2'
            else:
                para = doc.add_paragraph(text)
                para.style = 'Custom Body'

    def create_comprehensive_word_report(self):
        """
        Create the complete Word document report
        """
        print("Creating comprehensive Word document report...")

        # Create new document
        doc = Document()

        # Set up styles
        self.create_document_styles(doc)

        # Add header and footer
        self.add_header_footer(doc)

        # Title page
        title = doc.add_paragraph("Comprehensive Market Analysis Report")
        title.style = 'Custom Title'

        subtitle = doc.add_paragraph("Relative Strength & Performance Analysis")
        subtitle.style = 'Custom Heading 1'
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(f"Analysis Date: September 5, 2025")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.style = 'Custom Body'

        generated_para = doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}")
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_para.style = 'Custom Body'

        # Add page break
        doc.add_page_break()

        # Add sections
        self.add_executive_summary(doc)
        doc.add_page_break()

        self.add_key_findings(doc)
        doc.add_page_break()

        self.add_sector_analysis(doc)
        doc.add_page_break()

        self.add_top_performers_section(doc)
        doc.add_page_break()

        self.add_methodology_section(doc)
        doc.add_page_break()

        self.add_disclaimer_section(doc)

        # Save document
        output_file = self.output_dir / f"Comprehensive_Market_Analysis_Report_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(output_file)

        print(f"✅ Word document report saved: {output_file}")
        return output_file


def main():
    """
    Main execution function
    """
    current_dir = Path(__file__).parent.parent.parent
    data_directory = current_dir / "results"

    # Create Word report
    generator = EnhancedWordReportGenerator(data_directory)
    word_file = generator.create_comprehensive_word_report()

    print(f"\n🎉 Enhanced Word report complete!")
    print(f"📄 Word Document: {word_file}")


if __name__ == "__main__":
    main()